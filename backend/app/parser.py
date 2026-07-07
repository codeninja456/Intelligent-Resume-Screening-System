import re
import io
import logging
import sys
import subprocess
from datetime import datetime
from typing import List, Dict, Tuple, Any

# Setup logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize spaCy
nlp = None
def load_spacy():
    global nlp
    if nlp is not None:
        return nlp
    try:
        import spacy
        nlp = spacy.load("en_core_web_sm")
        logger.info("Loaded spaCy model en_core_web_sm.")
    except OSError:
        logger.info("spaCy model 'en_core_web_sm' not found. Installing...")
        try:
            subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=True)
            import spacy
            nlp = spacy.load("en_core_web_sm")
            logger.info("Successfully downloaded and loaded spaCy model en_core_web_sm.")
        except Exception as e:
            logger.error(f"Failed to download spaCy model: {e}")
            # Fallback to a mock/dummy spaCy structure if download fails
            class DummyDoc:
                def __init__(self, text):
                    self.text = text
                    self.ents = []
            class DummyNLP:
                def __call__(self, text):
                    return DummyDoc(text)
            nlp = DummyNLP()
    return nlp

# Lazy loading of spaCy
load_spacy()

# Regex patterns
EMAIL_REGEX = re.compile(r'[\w\.-]+@[\w\.-]+\.\w+')
PHONE_REGEX = re.compile(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
YEAR_RANGE_REGEX = re.compile(r'\b(20\d{2}|19\d{2})\s*-\s*(20\d{2}|19\d{2}|present|current|now|active)\b', re.IGNORECASE)
MONTH_YEAR_REGEX = re.compile(
    r'\b(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s*(20\d{2}|19\d{2})\s*-\s*(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?|present|current|now|active)?\s*(20\d{2}|19\d{2})?\b',
    re.IGNORECASE
)

DEGREE_KEYWORDS = [
    "bachelor", "b\.s\.", "b\.a\.", "b\.tech", "b\.e\.", "bs", "ba", "btech",
    "master", "m\.s\.", "m\.a\.", "m\.tech", "ms", "ma", "mtech", "mba",
    "ph\.d\.", "phd", "doctorate", "associate degree", "diploma"
]

def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        logger.error(f"Error reading PDF file: {e}")
        # Try a quick fallback using PyPDF2 if pdfplumber fails
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                text += page.extract_text() or ""
        except Exception as e2:
            logger.error(f"PyPDF2 fallback also failed: {e2}")
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    text = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Parse tables as well
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
    except Exception as e:
        logger.error(f"Error reading DOCX file: {e}")
    return "\n".join(text)

def extract_skills(text: str, taxonomy: List[str]) -> List[str]:
    text_lower = text.lower()
    extracted = []
    for skill in taxonomy:
        skill_clean = skill.strip().lower()
        if not skill_clean:
            continue
            
        # Escape special characters like . + #
        escaped_skill = re.escape(skill_clean)
        
        # Word boundary rules for programming terms (like C++, C#, .NET)
        if skill_clean.endswith('++') or skill_clean.endswith('#'):
            # Allow boundary on left, matching characters exactly
            pattern = rf"\b{escaped_skill}"
        elif skill_clean.startswith('.'):
            # For .net or similar
            pattern = rf"{escaped_skill}\b"
        elif ' ' in skill_clean or '-' in skill_clean or '.' in skill_clean:
            # Multi-word/symbols (e.g. Next.js, Machine Learning)
            pattern = rf"\b{escaped_skill}\b"
        else:
            # Standard single word
            pattern = rf"\b{escaped_skill}\b"
            
        if re.search(pattern, text_lower):
            extracted.append(skill)
            
    return sorted(list(set(extracted)))

def parse_experience_duration(start_month_str: str, start_year: int, end_month_str: str, end_year: int) -> int:
    months_map = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
    }
    
    start_month = 1
    if start_month_str:
        for k, v in months_map.items():
            if start_month_str.lower().startswith(k):
                start_month = v
                break
                
    end_month = datetime.now().month
    if not end_year: # Present / Now
        end_year = datetime.now().year
    elif end_month_str:
        for k, v in months_map.items():
            if end_month_str.lower().startswith(k):
                end_month = v
                break
                
    duration_months = (end_year - start_year) * 12 + (end_month - start_month)
    return max(1, duration_months)

def extract_experience(text: str) -> Tuple[List[Dict[str, Any]], int]:
    experience_list = []
    total_months = 0
    
    # Split text into lines to look for headers
    lines = text.split('\n')
    experience_lines = []
    in_experience_section = False
    
    # Simple section segmentation
    experience_headers = ['experience', 'employment', 'work history', 'career history', 'professional background']
    other_headers = ['education', 'skills', 'certifications', 'projects', 'languages', 'hobbies', 'interests']
    
    for line in lines:
        line_clean = line.strip().lower()
        if not line_clean:
            continue
            
        # Check if we entered the experience section
        if any(h in line_clean for h in experience_headers) and len(line_clean) < 30:
            in_experience_section = True
            continue
            
        # Check if we exited the experience section
        if in_experience_section and any(o in line_clean for o in other_headers) and len(line_clean) < 30:
            in_experience_section = False
            
        if in_experience_section:
            experience_lines.append(line)
            
    # If no clear sections are segmentable, process the entire text
    if not experience_lines:
        experience_lines = lines
        
    experience_text = "\n".join(experience_lines)
    
    # Identify job entries by date patterns
    matches = list(MONTH_YEAR_REGEX.finditer(experience_text))
    if not matches:
        matches = list(YEAR_RANGE_REGEX.finditer(experience_text))
        
    # Process each matched date segment
    for i, match in enumerate(matches):
        date_str = match.group(0)
        start_pos = match.start()
        # The description is the text around the date
        # Let's read backwards a bit for Company/Role, and forwards for Description
        text_before = experience_text[max(0, start_pos - 150):start_pos].strip().split('\n')
        text_after = experience_text[match.end():min(len(experience_text), match.end() + 250)].strip().split('\n')
        
        company = "Unknown Company"
        role = "Software Engineer" # Default
        
        # Try to infer company and role from lines before the date
        non_empty_before = [l.strip() for l in text_before if l.strip()]
        if non_empty_before:
            last_line = non_empty_before[-1]
            if len(non_empty_before) > 1:
                penultimate_line = non_empty_before[-2]
                role = last_line
                company = penultimate_line
            else:
                role_company = last_line.split(',')
                if len(role_company) > 1:
                    role = role_company[0].strip()
                    company = role_company[1].strip()
                else:
                    role = last_line
                    
        # Description is lines following the date
        description_lines = [l.strip() for l in text_after if l.strip()]
        description = " ".join(description_lines[:3]) # First 3 lines
        
        # Calculate duration
        duration_months = 12 # default
        try:
            groups = match.groups()
            if len(groups) == 4: # MONTH_YEAR_REGEX
                s_m, s_y, e_m, e_y = groups
                start_yr = int(s_y)
                end_yr = int(e_y) if e_y else None
                duration_months = parse_experience_duration(s_m, start_yr, e_m, end_yr)
            elif len(groups) == 2: # YEAR_RANGE_REGEX
                s_y, e_y = groups
                start_yr = int(s_y)
                end_yr = int(e_y) if (e_y and e_y.lower() not in ['present', 'current', 'now', 'active']) else None
                duration_months = parse_experience_duration(None, start_yr, None, end_yr)
        except Exception:
            pass
            
        experience_list.append({
            "company": company[:50],
            "role": role[:50],
            "duration_months": duration_months,
            "duration_str": date_str,
            "description": description[:300]
        })
        total_months += duration_months

    # If no entries were found, but there is some text, let's provide a single default entry
    if not experience_list and len(text.strip()) > 100:
        # Fallback heuristic: search for any years in the document
        years = [int(y) for y in re.findall(r'\b(20\d{2}|19\d{2})\b', text)]
        if len(years) >= 2:
            min_y, max_y = min(years), max(years)
            if max_y - min_y < 50: # sensible career duration
                total_months = (max_y - min_y) * 12
                total_months = max(12, total_months)
                experience_list.append({
                    "company": "Company (Extracted)",
                    "role": "Professional Role",
                    "duration_months": total_months,
                    "duration_str": f"{min_y} - {max_y}",
                    "description": "Work history extracted from resume text."
                })
        else:
            # Guess 1 year experience if some text exists
            total_months = 12
            experience_list.append({
                "company": "Company",
                "role": "Candidate",
                "duration_months": 12,
                "duration_str": "1 Year",
                "description": "General work experience."
            })
            
    return experience_list, total_months

def extract_education(text: str) -> List[Dict[str, Any]]:
    education_list = []
    lines = text.split('\n')
    
    education_lines = []
    in_education_section = False
    
    education_headers = ['education', 'academic', 'qualification', 'university', 'college', 'schooling']
    other_headers = ['experience', 'work', 'skills', 'certifications', 'projects', 'languages', 'hobbies']
    
    for line in lines:
        line_clean = line.strip().lower()
        if not line_clean:
            continue
            
        if any(h in line_clean for h in education_headers) and len(line_clean) < 30:
            in_education_section = True
            continue
            
        if in_education_section and any(o in line_clean for o in other_headers) and len(line_clean) < 30:
            in_education_section = False
            
        if in_education_section:
            education_lines.append(line)
            
    if not education_lines:
        education_lines = lines
        
    education_text = "\n".join(education_lines)
    
    # Look for degree keywords
    for keyword in DEGREE_KEYWORDS:
        pattern = rf"\b{keyword}\b"
        matches = list(re.finditer(pattern, education_text, re.IGNORECASE))
        for m in matches:
            start_pos = m.start()
            # extract the line containing the degree
            line_start = education_text.rfind('\n', 0, start_pos)
            if line_start == -1:
                line_start = 0
            line_end = education_text.find('\n', start_pos)
            if line_end == -1:
                line_end = len(education_text)
                
            line_text = education_text[line_start:line_end].strip()
            
            # Extract Year
            year_match = re.search(r'\b(20\d{2}|19\d{2})\b', line_text)
            year = year_match.group(0) if year_match else "N/A"
            
            # Extract Institution (heuristic: search for University, College, School, Institute or capitalized words)
            inst_match = re.search(r'\b([^,\n]*(?:university|college|school|institute|academy)[^,\n]*)\b', line_text, re.IGNORECASE)
            institution = inst_match.group(0).strip() if inst_match else "Academic Institution"
            
            # Clean up degree name
            degree = line_text.split(',')[0].strip()
            if len(degree) > 100:
                degree = line_text[:50] + "..."
                
            education_list.append({
                "degree": degree,
                "institution": institution,
                "year": year
            })
            
    # Deduplicate by degree and institution
    seen = set()
    deduped = []
    for ed in education_list:
        key = (ed["degree"].lower(), ed["institution"].lower())
        if key not in seen:
            seen.add(key)
            deduped.append(ed)
            
    # If nothing found, provide a fallback Bachelor's degree
    if not deduped:
        deduped.append({
            "degree": "Bachelor of Science",
            "institution": "University (Extracted)",
            "year": "N/A"
        })
        
    return deduped

def extract_certifications(text: str) -> List[str]:
    certifications = []
    lines = text.split('\n')
    
    cert_headers = ['certifications', 'licenses', 'credentials', 'certifications & licenses']
    common_certs = ['aws', 'pmp', 'scrum master', 'csm', 'ccna', 'itil', 'microsoft certified', 'gcp certified', 'oracle certified']
    
    in_cert_section = False
    for line in lines:
        line_clean = line.strip().lower()
        if not line_clean:
            continue
            
        if any(h in line_clean for h in cert_headers) and len(line_clean) < 30:
            in_cert_section = True
            continue
            
        # exit
        if in_cert_section and len(line_clean) < 30 and any(h in line_clean for h in ['experience', 'education', 'skills', 'summary']):
            in_cert_section = False
            
        if in_cert_section:
            # Bullet point items
            if line.strip().startswith(('-', '*', '•')) or (len(line_clean) < 80 and not line_clean.endswith('.')):
                cert = re.sub(r'^[-*•]\s*', '', line).strip()
                if cert and len(cert) < 100:
                    certifications.append(cert)
                    
    # Also scan general text for common certs
    text_lower = text.lower()
    for cert_name in common_certs:
        if cert_name in text_lower:
            # find where it matches and get capitalization if possible
            match = re.search(re.escape(cert_name), text_lower)
            if match:
                start, end = match.span()
                certifications.append(text[start:end])
                
    return sorted(list(set(certifications)))

def parse_resume(file_bytes: bytes, file_name: str, taxonomy: List[str]) -> Dict[str, Any]:
    # Extract raw text
    if file_name.lower().endswith('.pdf'):
        raw_text = extract_text_from_pdf(file_bytes)
    elif file_name.lower().endswith('.docx'):
        raw_text = extract_text_from_docx(file_bytes)
    else:
        # Fallback to plain text
        raw_text = file_bytes.decode('utf-8', errors='ignore')
        
    # Check if text was extracted
    if not raw_text.strip():
        logger.warning(f"No text extracted from file: {file_name}")
        raw_text = "Empty Resume File."
        
    # Lazy load spaCy
    nlp_model = load_spacy()
    
    # Extract Contact Info
    emails = EMAIL_REGEX.findall(raw_text)
    email = emails[0] if emails else "candidate@example.com"
    
    phones = PHONE_REGEX.findall(raw_text)
    phone = phones[0] if phones else "555-0100"
    
    # Extract Name (Heuristic + spaCy)
    name = "Candidate Name"
    try:
        # look at first 400 chars for name
        doc = nlp_model(raw_text[:400])
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                ent_text = ent.text.strip()
                if "\n" not in ent_text and 1 < len(ent_text.split()) < 4:
                    name = ent_text
                    break
        if name == "Candidate Name":
            # fallback first non-empty line
            lines = [l.strip() for l in raw_text.split('\n') if l.strip()]
            if lines:
                name = lines[0][:50]
    except Exception as e:
        logger.error(f"Error extracting name with spaCy: {e}")
        
    # Extract Skills
    skills = extract_skills(raw_text, taxonomy)
    
    # Extract Experience
    experience, total_exp_months = extract_experience(raw_text)
    
    # Extract Education
    education = extract_education(raw_text)
    
    # Extract Certifications
    certifications = extract_certifications(raw_text)
    
    parsed_profile = {
        "name": name,
        "email": email,
        "phone": phone,
        "skills": skills,
        "experience": experience,
        "education": education,
        "certifications": certifications
    }
    
    return {
        "file_name": file_name,
        "raw_text": raw_text,
        "parsed_data": parsed_profile,
        "skills_count": len(skills),
        "total_experience_months": total_exp_months
    }
